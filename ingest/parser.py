"""
文档解析器 - 从各种格式中提取纯文本
支持: PDF, DOCX, TXT, MD, XLSX, CSV
"""
from __future__ import annotations
from pathlib import Path
from dataclasses import dataclass
from typing import Optional
import re

from .loader import RawDocument


@dataclass
class ParsedDocument:
    doc_id: str
    category: str
    filename: str
    source_path: str
    raw_text: str          # 提取的原始纯文本
    page_count: int = 0
    parse_error: Optional[str] = None


class DocumentParser:
    """根据文件后缀分发解析逻辑"""

    def parse(self, raw_doc: RawDocument) -> ParsedDocument:
        ext = raw_doc.extension
        try:
            if ext == ".pdf":
                text, pages = self._parse_pdf(raw_doc.path)
            elif ext in (".docx", ".doc"):
                text, pages = self._parse_docx(raw_doc.path)
            elif ext in (".txt", ".md"):
                text, pages = self._parse_text(raw_doc.path)
            elif ext in (".xlsx", ".xls"):
                text, pages = self._parse_excel(raw_doc.path)
            elif ext == ".csv":
                text, pages = self._parse_csv(raw_doc.path)
            else:
                text, pages = "", 0
        except Exception as e:
            return ParsedDocument(
                doc_id=raw_doc.doc_id,
                category=raw_doc.category,
                filename=raw_doc.filename,
                source_path=str(raw_doc.path),
                raw_text="",
                parse_error=str(e),
            )
        return ParsedDocument(
            doc_id=raw_doc.doc_id,
            category=raw_doc.category,
            filename=raw_doc.filename,
            source_path=str(raw_doc.path),
            raw_text=text,
            page_count=pages,
        )

    # ── PDF ──────────────────────────────────────────────────────
    def _parse_pdf(self, fp: Path) -> tuple[str, int]:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(fp) as pdf:
            for page in pdf.pages:
                # 优先提取表格（信贷文件表格内容多）
                tables = page.extract_tables()
                table_text = ""
                for table in tables:
                    for row in table:
                        row_clean = [str(c).strip() if c else "" for c in row]
                        table_text += " | ".join(row_clean) + "\n"
                # 提取正文（去掉表格区域避免重复）
                body = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                combined = (body + "\n" + table_text).strip()
                if combined:
                    pages_text.append(combined)
        return "\n\n".join(pages_text), len(pages_text)

    # ── DOCX ─────────────────────────────────────────────────────
    def _parse_docx(self, fp: Path) -> tuple[str, int]:
        from docx import Document
        doc = Document(fp)
        parts = []
        # 段落
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        # 表格
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts), 1

    # ── TXT / MD ─────────────────────────────────────────────────
    def _parse_text(self, fp: Path) -> tuple[str, int]:
        for enc in ("utf-8", "gbk", "gb2312", "utf-16"):
            try:
                text = fp.read_text(encoding=enc)
                return text, 1
            except UnicodeDecodeError:
                continue
        return fp.read_text(errors="ignore"), 1

    # ── XLSX ─────────────────────────────────────────────────────
    def _parse_excel(self, fp: Path) -> tuple[str, int]:
        import pandas as pd
        xf = pd.ExcelFile(fp)
        parts = []
        for sheet in xf.sheet_names:
            df = xf.parse(sheet).fillna("")
            parts.append(f"【表格: {sheet}】")
            # 表头
            parts.append(" | ".join(str(c) for c in df.columns))
            # 数据行
            for _, row in df.iterrows():
                parts.append(" | ".join(str(v) for v in row.values))
        return "\n".join(parts), len(xf.sheet_names)

    # ── CSV ──────────────────────────────────────────────────────
    def _parse_csv(self, fp: Path) -> tuple[str, int]:
        import pandas as pd
        for enc in ("utf-8", "gbk", "gb2312"):
            try:
                df = pd.read_csv(fp, encoding=enc).fillna("")
                lines = [" | ".join(str(c) for c in df.columns)]
                for _, row in df.iterrows():
                    lines.append(" | ".join(str(v) for v in row.values))
                return "\n".join(lines), 1
            except Exception:
                continue
        return "", 0
